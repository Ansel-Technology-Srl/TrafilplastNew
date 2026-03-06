#!/usr/bin/env python3
"""Genera Analisi Funzionale DOCX per Recinzioni Portal."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DIAGRAMS = '/home/user/TrafilplastNew/diagrams'
OUT = '/home/user/TrafilplastNew/ANALISI_FUNZIONALE.docx'

doc = Document()

# ─── Stili ──────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    hs.font.name = 'Calibri'

def add_styled_table(doc, headers, rows, col_widths=None):
    """Crea tabella con header colorato."""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1e40af"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                r.font.bold = True
                r.font.size = Pt(10)
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table

def add_image(doc, path, width=Inches(6)):
    """Aggiunge immagine centrata."""
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)

# ═══════════════════════════════════════════════════════════════
# COPERTINA
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ANALISI FUNZIONALE')
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Recinzioni Portal - Portale B2B')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('Trafilplast S.r.l.')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
run.bold = True

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Versione 1.0 - Marzo 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Documento di analisi funzionale completa del portale web B2B\nper la gestione del catalogo prodotti, configurazione recinzioni,\npreventivi, ordini e trasmissione EDI EURITMO.')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# INDICE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Indice', level=1)
indice = [
    '1. Introduzione e Obiettivi',
    '2. Architettura del Sistema',
    '3. Profili Utente e Permessi',
    '4. Flusso Generale dell\'Applicazione',
    '5. Modulo Autenticazione',
    '6. Modulo Catalogo Prodotti',
    '7. Modulo Configuratore Recinzioni 3D',
    '8. Modulo Carrello',
    '9. Modulo Preventivi e Ordini',
    '10. Modulo Gestione Listini (Super User)',
    '11. Modulo Amministrazione',
    '12. Funzionalita Trasversali',
    '13. Specifiche API REST',
    '14. Requisiti Non Funzionali',
]
for item in indice:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUZIONE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Introduzione e Obiettivi', level=1)

doc.add_heading('1.1 Scopo del Documento', level=2)
doc.add_paragraph(
    'Il presente documento descrive in modo dettagliato l\'analisi funzionale del portale web '
    '"Recinzioni Portal", un\'applicazione B2B sviluppata per Trafilplast S.r.l. destinata a '
    'rivenditori e dealer del settore recinzioni. Il documento copre tutti i flussi operativi, '
    'le interfacce utente, le logiche di business e le integrazioni del sistema.'
)

doc.add_heading('1.2 Obiettivi del Portale', level=2)
objectives = [
    'Catalogo prodotti: Consultazione completa del catalogo recinzioni con filtri avanzati, ricerca testuale, visualizzazione griglia/lista e dettagli tecnici multilingua.',
    'Configuratore 3D: Strumento interattivo per la composizione di recinzioni personalizzate con visualizzazione 2D/3D in tempo reale, calcolo automatico della distinta base (BOM) e prezzi.',
    'Gestione preventivi e ordini: Flusso completo dalla composizione del carrello alla generazione del preventivo, conferma ordine e trasmissione al fornitore.',
    'Trasmissione EDI: Generazione automatica di file EURITMO ORDERS (standard GS1 Italy, Release 25.1) per l\'invio elettronico degli ordini al fornitore.',
    'Gestione listini: Creazione e gestione di listini prezzi personalizzati per punto vendita con confronto rispetto al listino pubblico.',
    'Amministrazione: Gestione utenti, import/export massivo dati (prodotti, prezzi, clienti, punti vendita) tramite Excel/CSV.',
    'Multilingua: Supporto completo per 4 lingue (Italiano, English, Francais, Deutsch).',
    'Accessibilita: Conformita WCAG 2.1 livello AA per garantire l\'accesso a tutti gli utenti.',
    'PWA: Funzionamento come Progressive Web App con supporto offline e installabilita.',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading('1.3 Utenti Target', level=2)
doc.add_paragraph(
    'Il portale e destinato a 4 categorie di utenti appartenenti alla rete distributiva Trafilplast:'
)
add_styled_table(doc,
    ['Tipo', 'Codice', 'Descrizione', 'Accesso Principale'],
    [
        ['Amministratore', '1', 'Gestione completa del sistema', 'Utenti, Import/Export dati'],
        ['Super User', '2', 'Responsabile listini cliente', 'Gestione listini prezzi personalizzati'],
        ['Capo Negozio', '3', 'Manager punto vendita', 'Catalogo, Configuratore, Ordini (tutto il PdV)'],
        ['Operatore', '4', 'Operatore punto vendita', 'Catalogo, Configuratore, Ordini (solo propri)'],
    ],
    [3, 2, 5, 6]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. ARCHITETTURA
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Architettura del Sistema', level=1)

doc.add_heading('2.1 Architettura Generale', level=2)
doc.add_paragraph(
    'L\'applicazione segue un\'architettura client-server moderna con separazione netta tra frontend e backend:'
)
add_image(doc, f'{DIAGRAMS}/08_architettura.png', Inches(6.2))

doc.add_heading('2.2 Stack Tecnologico', level=2)
doc.add_heading('Frontend', level=3)
add_styled_table(doc,
    ['Tecnologia', 'Versione', 'Utilizzo'],
    [
        ['React', '18.3.1', 'Framework UI con componenti funzionali e hooks'],
        ['Vite', '6.0.0', 'Build tool, dev server con HMR, chunking ottimizzato'],
        ['React Router', '6.28.0', 'Routing SPA con layout nested e route protette'],
        ['Zustand', '5.0.0', 'State management (auth store + cart store)'],
        ['Three.js', '0.170.0', 'Rendering 3D del configuratore recinzioni'],
        ['React Three Fiber', '-', 'Integrazione React per Three.js'],
        ['Tailwind CSS', '3.4.15', 'Utility-first CSS framework con tema custom'],
        ['i18next', '24.0.0', 'Internazionalizzazione (IT, EN, FR, DE)'],
        ['Lucide React', '0.460.0', 'Libreria icone SVG'],
        ['vite-plugin-pwa', '0.21.0', 'Progressive Web App con Workbox'],
        ['react-hot-toast', '2.4.1', 'Notifiche toast'],
    ],
    [4, 2, 10]
)

doc.add_paragraph()
doc.add_heading('Backend', level=3)
add_styled_table(doc,
    ['Tecnologia', 'Versione', 'Utilizzo'],
    [
        ['ASP.NET Core', '8.0', 'Framework Web API con controller-based routing'],
        ['Entity Framework Core', '8.0.11', 'ORM per SQL Server'],
        ['SQL Server', '-', 'Database relazionale (Latin1_General_CI_AS)'],
        ['JWT Bearer', 'HS256', 'Autenticazione con token + refresh token'],
        ['BCrypt', '-', 'Hashing password (cost factor 12)'],
        ['QuestPDF', '2026.2.1', 'Generazione documenti PDF'],
        ['ClosedXML', '0.104.2', 'Import/Export Excel'],
        ['SMTP', '-', 'Invio email (ordini, reset password)'],
    ],
    [4, 2, 10]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. PROFILI UTENTE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Profili Utente e Permessi', level=1)

doc.add_heading('3.1 Matrice di Accesso', level=2)
doc.add_paragraph(
    'Ogni utente e associato a un tipo (1-4) che determina le funzionalita accessibili. '
    'Gli utenti di tipo 3 e 4 sono anche associati a un Cliente (ItemID) e a un Punto di Vendita (ItemIDSede).'
)
add_styled_table(doc,
    ['Funzionalita', 'Admin (1)', 'Super User (2)', 'Manager (3)', 'Operatore (4)'],
    [
        ['Login/Logout', 'Si', 'Si', 'Si', 'Si'],
        ['Cambio Password', 'Si', 'Si', 'Si', 'Si'],
        ['Gestione Utenti', 'Si', 'No', 'No', 'No'],
        ['Import/Export Dati', 'Si', 'No', 'No', 'No'],
        ['Gestione Listini', 'No', 'Si', 'No', 'No'],
        ['Catalogo Prodotti', 'No', 'No', 'Si', 'Si'],
        ['Configuratore 3D', 'No', 'No', 'Si', 'Si'],
        ['Carrello', 'No', 'No', 'Si', 'Si'],
        ['Preventivi/Ordini', 'No', 'No', 'Si (tutto PdV)', 'Si (solo propri)'],
        ['Invio EDI', 'No', 'No', 'Si', 'Si'],
        ['Download PDF', 'No', 'No', 'Si', 'Si'],
    ],
    [4, 2.5, 2.5, 2.5, 2.5]
)

doc.add_heading('3.2 Redirect Automatico per Ruolo', level=2)
doc.add_paragraph(
    'Dopo il login, l\'utente viene automaticamente reindirizzato alla pagina principale '
    'corrispondente al proprio ruolo:'
)
add_styled_table(doc,
    ['Tipo Utente', 'Pagina Home', 'Route'],
    [
        ['Amministratore (1)', 'Gestione Utenti', '/admin/utenti'],
        ['Super User (2)', 'Gestione Listini', '/listini'],
        ['Manager/Operatore (3-4)', 'Catalogo Prodotti', '/catalogo'],
    ],
    [4, 5, 5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. FLUSSO GENERALE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. Flusso Generale dell\'Applicazione', level=1)

doc.add_paragraph(
    'Il diagramma seguente illustra il flusso generale dell\'applicazione, dalla fase di login '
    'fino alle operazioni specifiche per ciascun profilo utente.'
)
add_image(doc, f'{DIAGRAMS}/01_flusso_generale.png', Inches(6.2))

doc.add_heading('4.1 Descrizione del Flusso Principale', level=2)
doc.add_paragraph(
    '1. L\'utente accede al portale tramite browser o PWA installata.\n'
    '2. Viene presentata la pagina di login con selezione lingua.\n'
    '3. Dopo autenticazione con successo, il sistema verifica il tipo utente.\n'
    '4. L\'utente viene reindirizzato alla sezione appropriata:\n'
    '   - Admin: Gestione utenti e import/export dati\n'
    '   - Super User: Gestione listini prezzi personalizzati\n'
    '   - Shop User (Manager/Operatore): Catalogo, Configuratore, Carrello, Ordini\n'
    '5. Per gli utenti Shop, il flusso principale prevede:\n'
    '   a. Navigazione catalogo e/o configurazione recinzione\n'
    '   b. Aggiunta prodotti al carrello\n'
    '   c. Compilazione preventivo con dati fatturazione/consegna\n'
    '   d. Conferma ordine\n'
    '   e. Trasmissione ordine al fornitore via EDI EURITMO'
)

doc.add_heading('4.2 Navigazione e Layout', level=2)
doc.add_paragraph(
    'L\'applicazione utilizza un layout a sidebar con navigazione laterale persistente (desktop) '
    'o collassabile (mobile). La sidebar contiene:'
)
nav_items = [
    'Logo e nome applicazione',
    'Nome utente e badge tipo utente',
    'Menu di navigazione contestuale al ruolo (con icone)',
    'Badge contatore carrello (per Shop User)',
    'Selettore lingua (IT, EN, FR, DE)',
    'Pulsante cambio password',
    'Pulsante logout',
]
for item in nav_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. MODULO AUTENTICAZIONE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. Modulo Autenticazione', level=1)

doc.add_heading('5.1 Diagramma di Flusso', level=2)
add_image(doc, f'{DIAGRAMS}/02_flusso_autenticazione.png', Inches(6.2))

doc.add_heading('5.2 Pagina di Login', level=2)
doc.add_paragraph('La pagina di login presenta i seguenti elementi:')
login_elements = [
    'Selettore lingua: 4 pulsanti (IT, EN, FR, DE) posizionati in alto a destra, permettono di cambiare la lingua dell\'interfaccia prima ancora di effettuare il login.',
    'Campo Username: Input testuale con icona utente, obbligatorio.',
    'Campo Password: Input password con pulsante toggle visibilita (mostra/nascondi), obbligatorio.',
    'Checkbox "Ricordami": Se attivato, il token viene salvato in localStorage (persistente); altrimenti in sessionStorage (sessione browser).',
    'Link "Password dimenticata": Naviga alla pagina di recupero password.',
    'Pulsante Login: Invia le credenziali all\'API. Mostra spinner durante il caricamento.',
    'Messaggio di errore: Appare sotto il form in caso di credenziali errate o account bloccato.',
]
for item in login_elements:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.3 Sicurezza Autenticazione', level=2)
doc.add_paragraph('Il sistema implementa le seguenti misure di sicurezza:')
security = [
    'Hashing BCrypt: Le password sono hashate con BCrypt (cost factor 12) nel database.',
    'Protezione brute-force: Dopo 5 tentativi falliti consecutivi, l\'account viene bloccato per 15 minuti. Il contatore si resetta dopo un login riuscito.',
    'Token JWT: Il token di accesso ha durata limitata e contiene claims per UserID, UserType, ItemID, ItemIDSede.',
    'Refresh Token: Token opaco salvato nel database con scadenza piu lunga, utilizzato per rinnovare il JWT senza re-login.',
    'Auto-refresh: Il client intercetta risposte 401 e tenta automaticamente il refresh del token prima di ripresentare la richiesta originale.',
    'Complessita password: Minimo 8 caratteri con almeno una maiuscola, una minuscola, un numero e un carattere speciale.',
]
for item in security:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.4 Recupero Password', level=2)
doc.add_paragraph(
    'Il flusso di recupero password si articola in tre fasi:\n\n'
    '1. Richiesta reset: L\'utente inserisce il proprio indirizzo email. Il sistema genera un token '
    'univoco con scadenza (24 ore) e invia un\'email con un link di reset. Per sicurezza, la risposta '
    'API e sempre positiva indipendentemente dall\'esistenza dell\'email.\n\n'
    '2. Validazione token: Quando l\'utente clicca il link nell\'email, il sistema verifica la validita '
    'del token (esistenza e scadenza). Se il token non e valido, viene mostrato un messaggio di errore '
    'con link per richiedere un nuovo reset.\n\n'
    '3. Nuova password: L\'utente inserisce la nuova password con conferma. Un indicatore visuale '
    'mostra la robustezza della password in tempo reale. Dopo il reset riuscito, l\'utente viene '
    'reindirizzato alla pagina di login.'
)

doc.add_heading('5.5 Cambio Password', level=2)
doc.add_paragraph(
    'Tutti gli utenti autenticati possono cambiare la propria password tramite la modale '
    'accessibile dalla sidebar. Il form richiede:'
)
pwd_fields = [
    'Password attuale (verifica identita)',
    'Nuova password (con indicatore di robustezza visuale)',
    'Conferma nuova password (con indicatore di corrispondenza)',
]
for item in pwd_fields:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. MODULO CATALOGO
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. Modulo Catalogo Prodotti', level=1)

doc.add_heading('6.1 Diagramma di Flusso', level=2)
add_image(doc, f'{DIAGRAMS}/03_flusso_catalogo.png', Inches(6.2))

doc.add_heading('6.2 Interfaccia Catalogo', level=2)
doc.add_paragraph(
    'La pagina catalogo e la pagina principale per gli utenti Shop (tipo 3-4) e presenta '
    'un\'interfaccia ricca per la consultazione e selezione dei prodotti.'
)

doc.add_heading('6.2.1 Barra di Ricerca e Filtri', level=3)
doc.add_paragraph(
    'L\'area superiore della pagina contiene:'
)
filter_items = [
    'Campo di ricerca testuale: Permette la ricerca per codice prodotto o descrizione. La ricerca viene eseguita alla pressione di Invio o click sul pulsante cerca.',
    'Filtro Categoria: Menu a tendina che mostra tutte le categorie disponibili con conteggio prodotti. La selezione di una categoria aggiorna automaticamente il filtro Famiglia.',
    'Filtro Famiglia: Menu a tendina che dipende dalla categoria selezionata. Mostra solo le famiglie della categoria corrente.',
    'Filtro Gruppo: Menu a tendina per ulteriore raffinamento.',
    'Ordinamento: Selettore per campo di ordinamento (codice, descrizione, categoria) con pulsante per direzione ascendente/discendente.',
    'Filtri attivi: I filtri applicati vengono mostrati come "pills" rimovibili sotto la barra. Ogni pill mostra il tipo di filtro e il valore, con pulsante X per rimuoverlo.',
]
for item in filter_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.2.2 Modalita di Visualizzazione', level=3)
doc.add_paragraph('Due pulsanti toggle permettono di scegliere la modalita:')
doc.add_paragraph(
    'Vista Griglia: Prodotti visualizzati come cards in una griglia responsive (6 colonne desktop, '
    '2 mobile). Ogni card mostra: immagine prodotto, codice, nome, tags (categoria/famiglia), '
    'prezzo unitario, pulsante azione (aggiungi al carrello o configura).',
    style='List Bullet'
)
doc.add_paragraph(
    'Vista Lista: Tabella con colonne: Codice, Nome, Categoria, Famiglia, UM, Prezzo, Azioni. '
    'Piu compatta, ideale per la consultazione rapida.',
    style='List Bullet'
)

doc.add_heading('6.2.3 Paginazione', level=3)
doc.add_paragraph(
    'I risultati sono paginati (24 prodotti per pagina in griglia). La barra di paginazione '
    'mostra il range corrente ("Mostrando 1-24 di 352 prodotti") e i pulsanti pagina con '
    'navigazione diretta, precedente e successiva.'
)

doc.add_heading('6.3 Dettaglio Prodotto', level=2)
doc.add_paragraph(
    'Cliccando su un prodotto si apre una modale overlay che presenta:'
)
detail_items = [
    'Header: Codice prodotto (monospace), nome completo, badge UM, badge "Configurabile" se applicabile.',
    'Sezione Prezzo: Prezzo unitario formattato con valuta, codice listino, pulsante azione principale (Configura / Aggiungi al Carrello).',
    'Informazioni Tecniche: Griglia con icone che mostra: posizione archivio, origine, sede, gruppo, categoria, albero, famiglia, distinta base, tipo configurazione, colore configurazione.',
    'Immagine Prodotto: Visualizzazione dell\'immagine del prodotto (200x200px) se disponibile.',
    'Traduzioni: Sezione con le traduzioni del prodotto in tutte le lingue disponibili (IT come riferimento, poi EN, FR, DE). Ogni lingua mostrata con badge colorato.',
]
for item in detail_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.4 Risoluzione Prezzi', level=2)
doc.add_paragraph(
    'I prezzi mostrati nel catalogo seguono una cascata di priorita per identificare il prezzo '
    'corretto per l\'utente corrente:'
)
doc.add_paragraph(
    '1. Listino personalizzato del punto vendita (se esiste e ha date di validita correnti)\n'
    '2. Listino di default del punto vendita (LstCod del PdV)\n'
    '3. Listino pubblico (LstCodPubb)\n'
    '4. Listino di default del cliente'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. CONFIGURATORE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. Modulo Configuratore Recinzioni 3D', level=1)

doc.add_heading('7.1 Diagramma di Flusso', level=2)
add_image(doc, f'{DIAGRAMS}/04_flusso_configuratore.png', Inches(6.2))

doc.add_heading('7.2 Panoramica', level=2)
doc.add_paragraph(
    'Il Configuratore Recinzioni e lo strumento piu avanzato del portale. Permette agli utenti '
    'di comporre recinzioni personalizzate selezionando parametri specifici e visualizzando il '
    'risultato in tempo reale sia in 2D che in 3D. Il processo e guidato tramite uno stepper a 3 fasi.'
)

doc.add_heading('7.3 Step 1: Parametri Base', level=2)
doc.add_paragraph('Nel primo step l\'utente seleziona i parametri fondamentali della recinzione:')

doc.add_heading('Colore', level=3)
doc.add_paragraph(
    'L\'utente puo scegliere tra colori predefiniti (presentati come pulsanti colorati) o inserire '
    'un codice esadecimale personalizzato. Una checkbox "Stesso colore" permette di usare lo stesso '
    'colore per doghe e pali; disattivandola, appaiono due selettori separati per colore doghe e colore pali.'
)

doc.add_heading('Tipo di Fissaggio', level=3)
doc.add_paragraph(
    'Due pulsanti illustrati permettono di scegliere tra:\n'
    '- Cemento: Pali da tassellare su pavimentazione esistente\n'
    '- Terreno: Pali da interrare direttamente nel suolo'
)

doc.add_heading('Tipo Doghe', level=3)
doc.add_paragraph(
    'Due pulsanti con descrizione:\n'
    '- Persiana: Doghe inclinate tipo persiana (ventilazione)\n'
    '- Pieno: Doghe piene (privacy totale)'
)

doc.add_heading('Altezza', level=3)
doc.add_paragraph(
    'Quattro opzioni di altezza: 100 cm, 150 cm, 185 cm, 200 cm. Ogni pulsante mostra anche '
    'il numero di doghe risultante per quella altezza. L\'altezza influenza direttamente il '
    'calcolo dei componenti nella distinta base.'
)

doc.add_heading('7.4 Step 2: Design Sezioni', level=2)
doc.add_paragraph(
    'Nel secondo step l\'utente definisce la geometria della recinzione aggiungendo sezioni:'
)
design_items = [
    'Lista sezioni: Ogni sezione e un pannello espandibile che mostra lunghezza e angolo.',
    'Lunghezza: Slider + campo numerico (range 10-158 cm). Lo slider fornisce feedback visuale immediato.',
    'Angolo: Checkbox per angolo a 90 gradi. Se attivato, la sezione successiva forma un angolo retto.',
    'Aggiungi sezione: Pulsante per aggiungere una nuova sezione (massimo 20 sezioni).',
    'Rimuovi sezione: Ogni sezione (tranne la prima) ha un pulsante di eliminazione.',
    'Riepilogo rapido: Box informativo che mostra numero sezioni, numero pali, lunghezza totale.',
]
for item in design_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7.5 Visualizzazione 2D/3D', level=2)
doc.add_paragraph(
    'Il pannello destro del configuratore mostra l\'anteprima della recinzione con 3 modalita:'
)
doc.add_paragraph(
    '- Vista 3D: Scena Three.js interattiva con rotazione, zoom e pan. Mostra la recinzione '
    'con materiali e colori realistici. I pali sono cliccabili per la selezione.\n'
    '- Vista 2D: Pianta schematica dall\'alto con quote e dimensioni. Evidenzia angoli e sezioni.\n'
    '- Vista Combinata: Entrambe le viste affiancate (split view).\n\n'
    'Un pulsante fullscreen permette di espandere la vista a tutto schermo. '
    'Per l\'accessibilita, e disponibile un componente AccessibleConfigurator con '
    'descrizione testuale alternativa della configurazione.'
)

doc.add_heading('7.6 Step 3: Riepilogo e Distinta Base', level=2)
doc.add_paragraph(
    'Nel terzo step il sistema calcola la distinta base (BOM) inviando i parametri all\'API '
    'POST /api/configuratore/distinta-base. Il risultato include:'
)
bom_items = [
    'Riepilogo configurazione: Numero sezioni, numero pali, doghe per sezione, lunghezza totale.',
    'Lista componenti: Ogni componente mostra codice prodotto, descrizione, unita di misura, quantita, prezzo unitario e prezzo totale. I componenti sono espandibili/collassabili.',
    'Totale complessivo: Somma di tutti i componenti, formattata con valuta.',
    'Pulsante "Aggiungi al Carrello": Aggiunge l\'intera configurazione come un singolo articolo al carrello, mantenendo tutti i parametri e la distinta base associata.',
]
for item in bom_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7.7 Componenti Calcolati', level=2)
doc.add_paragraph('Il servizio backend calcola i seguenti componenti in base ai parametri:')
add_styled_table(doc,
    ['Componente', 'Logica di Calcolo'],
    [
        ['Pali (Montanti)', 'N. sezioni + 1. Tipo varia per altezza e fissaggio (cemento/terreno)'],
        ['Doghe', 'Per ogni sezione: N. doghe dipende da altezza. Tipo varia per stile (persiana/pieno)'],
        ['Fissaggi', 'Set di viti/tasselli per ogni palo. Tipo varia per fissaggio'],
        ['Distanziali', 'Tra ogni doga, per ogni sezione'],
        ['Cappucci', 'Uno per ogni palo (copertura superiore)'],
        ['Giunti angolari', 'Per ogni sezione con angolo a 90 gradi'],
    ],
    [4, 12]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. MODULO CARRELLO
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. Modulo Carrello', level=1)

doc.add_heading('8.1 Interfaccia Carrello', level=2)
doc.add_paragraph(
    'Il carrello gestisce temporaneamente i prodotti selezionati dall\'utente prima della '
    'creazione del preventivo. Lo stato e mantenuto in memoria tramite Zustand store e non '
    'persiste tra sessioni browser.'
)

doc.add_heading('8.2 Elementi dell\'Interfaccia', level=2)
cart_items = [
    'Header: Titolo con icona carrello, badge con numero articoli, pulsante "Svuota carrello" (rosso).',
    'Stato vuoto: Se il carrello e vuoto, viene mostrata un\'icona grande, il messaggio "Il carrello e vuoto" e un pulsante per tornare al catalogo.',
    'Tabella articoli: Colonne - Espandi (per configurazioni), Prodotto (nome + codice + badge), UM, Quantita (con pulsanti +/-), Prezzo unitario, Totale, Azioni.',
    'Articoli semplici: La quantita e modificabile con pulsanti incremento/decremento. Il prezzo unitario e visibile.',
    'Articoli configurati: La quantita e fissa (impostata dal configuratore). Un badge "Configurato" identifica il prodotto. Espandendo la riga si vedono i parametri di configurazione e la lista componenti con prezzi individuali.',
    'Pulsante "Riconfigura": Per gli articoli configurati, permette di tornare al configuratore con i parametri precaricati per modificare la configurazione.',
    'Riepilogo: Card con subtotale, IVA (22% fisso), totale complessivo.',
    'Azioni finali: Pulsante "Continua acquisti" (torna al catalogo) e "Salva preventivo" (procede al form preventivo).',
]
for item in cart_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. MODULO PREVENTIVI E ORDINI
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. Modulo Preventivi e Ordini', level=1)

doc.add_heading('9.1 Diagramma di Flusso', level=2)
add_image(doc, f'{DIAGRAMS}/05_flusso_ordini.png', Inches(6.2))

doc.add_heading('9.2 Form Preventivo', level=2)
doc.add_paragraph(
    'Il form preventivo raccoglie i dati necessari per la creazione dell\'ordine. '
    'Il layout e organizzato in due colonne responsive:'
)

doc.add_heading('Dati Fatturazione (obbligatori)', level=3)
add_styled_table(doc,
    ['Campo', 'Tipo', 'Validazione', 'Descrizione'],
    [
        ['Ragione Sociale', 'Testo', 'Min 3 caratteri', 'Nome azienda/cliente'],
        ['Indirizzo', 'Testo', 'Obbligatorio', 'Via e numero civico'],
        ['CAP', 'Testo', '5 cifre', 'Codice Avviamento Postale'],
        ['Citta', 'Testo', 'Obbligatorio', 'Comune'],
        ['Provincia', 'Testo', '2 lettere maiuscole', 'Sigla provincia'],
        ['P.IVA', 'Testo', 'IT + 11 cifre', 'Partita IVA'],
        ['Codice Fiscale', 'Testo', '16 caratteri alfanumerici', 'Codice fiscale'],
    ],
    [3, 2, 4, 5]
)

doc.add_heading('Dati Consegna', level=3)
doc.add_paragraph(
    'Una checkbox "Uguale alla fatturazione" pre-compila e disabilita i campi consegna '
    'con i dati di fatturazione. Se deselezionata, i campi diventano editabili e obbligatori.'
)

doc.add_heading('Pagamento e Note', level=3)
doc.add_paragraph(
    'Campi opzionali per codice/descrizione pagamento e un\'area testo per note libere.'
)

doc.add_heading('Pre-compilazione', level=2)
doc.add_paragraph(
    'All\'apertura del form, il sistema carica automaticamente i dati anagrafici dell\'utente '
    'tramite GET /api/ordini/dati-anagrafica, pre-compilando i campi fatturazione e consegna '
    'con i dati del cliente e punto vendita associati.'
)

doc.add_heading('9.3 Gestione Ordini', level=2)
doc.add_paragraph(
    'La pagina ordini mostra l\'elenco di tutti i preventivi e ordini dell\'utente con le '
    'seguenti funzionalita:'
)

doc.add_heading('Filtri e Ricerca', level=3)
doc.add_paragraph(
    '- Pulsanti stato: "Tutti", "Preventivo", "Ordine", "Carrello" con conteggi. '
    'Ogni pulsante filtra la lista per lo stato corrispondente.\n'
    '- Campo ricerca: Ricerca per numero ordine o ragione sociale.'
)

doc.add_heading('Tabella Ordini', level=3)
add_styled_table(doc,
    ['Colonna', 'Contenuto', 'Note'],
    [
        ['Numero', '#123', 'Identificativo univoco auto-incrementale'],
        ['Data', 'GG/MM/AAAA', 'Data di creazione, formattata per locale'],
        ['Stato', 'Badge colorato', 'Grigio=Carrello, Blu=Preventivo, Verde=Ordine'],
        ['Ragione Sociale', 'Nome azienda', 'Dati fatturazione'],
        ['Totale', 'EUR formattato', 'Totale IVA inclusa'],
        ['Inviato', 'Icona check', 'Indica se trasmesso via EDI'],
        ['Azioni', 'Pulsanti', 'Dettaglio, PDF, Conferma, Invia, Elimina'],
    ],
    [3, 3, 8]
)

doc.add_heading('9.4 Dettaglio Ordine', level=2)
doc.add_paragraph(
    'Cliccando sull\'icona "dettaglio" si apre una modale overlay con:'
)
detail_items = [
    'Header con numero ordine, stato e pulsante chiudi.',
    'Sezione dati fatturazione e consegna (due colonne).',
    'Sezione pagamento e note (se presenti).',
    'Tabella righe ordine con: prodotto, codice, badge configurato, UM, quantita, prezzo unitario, totale.',
    'Per articoli configurati: riga espandibile con parametri configurazione e sotto-componenti.',
    'Totali: subtotale, IVA 22%, totale complessivo.',
    'Barra azioni: Download PDF, Conferma ordine, Preview EURITMO, Invia a fornitore, Elimina.',
]
for item in detail_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('9.5 Ciclo di Vita dell\'Ordine', level=2)
doc.add_paragraph(
    'Un ordine attraversa i seguenti stati:'
)
add_styled_table(doc,
    ['Stato', 'Badge', 'Azioni Disponibili', 'Descrizione'],
    [
        ['Preventivo', 'Blu', 'Modifica righe, Conferma, PDF, Elimina', 'Bozza modificabile'],
        ['Ordine', 'Verde', 'PDF, Preview EURITMO, Invia, Elimina righe', 'Confermato, pronto per invio'],
        ['Inviato', 'Verde + Check', 'PDF, Preview EURITMO, Re-invia', 'Trasmesso al fornitore via EDI'],
    ],
    [2.5, 2, 6, 4]
)

doc.add_heading('9.6 Trasmissione EURITMO EDI', level=2)
doc.add_paragraph(
    'La funzionalita di trasmissione EDI genera un file nel formato GS1 Italy EURITMO ORDERS '
    'Release 25.1. Il file e composto da record a lunghezza fissa con i seguenti tipi:'
)
add_styled_table(doc,
    ['Tipo Record', 'Contenuto'],
    [
        ['BGM', 'Intestazione messaggio (numero ordine, data, tipo documento)'],
        ['NAS', 'Fornitore (dati aziendali Trafilplast)'],
        ['CTA', 'Contatto (referente ordine)'],
        ['NAB', 'Acquirente (dati cliente)'],
        ['NAD', 'Destinazione merce (punto vendita)'],
        ['NAI', 'Indirizzo fatturazione'],
        ['DTM', 'Date (consegna richiesta, ordine)'],
        ['FTX', 'Testo libero (note ordine)'],
        ['PAT', 'Condizioni di pagamento'],
        ['LIN', 'Righe ordine (codice EAN, quantita, prezzo)'],
        ['CNT', 'Contatori finali (totale righe, totale quantita)'],
    ],
    [3, 13]
)

doc.add_paragraph(
    '\nLa modale di preview EURITMO mostra il contenuto del file EDI con colorazione '
    'per tipo di record. I pulsanti permettono di scaricare il file .edi e di inviarlo '
    'al fornitore (che riceve anche una notifica email con il file in allegato).'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. MODULO LISTINI
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. Modulo Gestione Listini (Super User)', level=1)

doc.add_heading('10.1 Diagramma di Flusso', level=2)
add_image(doc, f'{DIAGRAMS}/07_flusso_listini.png', Inches(6.2))

doc.add_heading('10.2 Panoramica', level=2)
doc.add_paragraph(
    'Il modulo listini e accessibile esclusivamente agli utenti Super User (tipo 2) e permette '
    'la gestione dei listini prezzi personalizzati per ciascun punto vendita del cliente assegnato.'
)

doc.add_heading('10.3 Interfaccia', level=2)
doc.add_paragraph('L\'interfaccia e divisa in due pannelli:')
doc.add_paragraph(
    'Pannello sinistro (40%): Lista dei punti vendita del cliente assegnato. Ogni voce mostra: '
    'nome negozio, localita, stato del listino personalizzato (attivo/scaduto/assente), '
    'numero di prodotti a prezzo personalizzato.',
    style='List Bullet'
)
doc.add_paragraph(
    'Pannello destro (60%): Contenuto contestuale basato sulla selezione e sullo stato del listino:',
    style='List Bullet'
)

doc.add_heading('10.3.1 Senza Listino Personalizzato', level=3)
doc.add_paragraph(
    'Viene mostrato il listino pubblico in modalita di sola lettura con un pulsante "Crea Listino '
    'Personalizzato". Cliccando si apre un form inline per specificare le date di validita '
    '(dal/al). Alla conferma, viene creato un nuovo listino personalizzato copiando tutti i '
    'prezzi dal listino pubblico.'
)

doc.add_heading('10.3.2 Con Listino Personalizzato', level=3)
doc.add_paragraph('L\'editor del listino personalizzato offre:')
list_features = [
    'Header: Codice listino, date di validita, pulsante modifica date.',
    'Statistiche: 3 card cliccabili (Totale prodotti, Modificati, Invariati) che fungono anche da filtro.',
    'Ricerca: Campo per cercare prodotti per codice o descrizione.',
    'Tabella confronto: Per ogni prodotto mostra codice, descrizione, UM, prezzo pubblico, prezzo personalizzato, differenza percentuale (con freccia su/giu).',
    'Modifica inline: Cliccando sul prezzo personalizzato si attiva l\'editing diretto. Enter per salvare, Esc per annullare. I prezzi modificati sono evidenziati in ambra.',
    'Reset prezzi: Pulsante per ripristinare tutti i prezzi ai valori del listino pubblico.',
    'Elimina listino: Pulsante per eliminare completamente il listino personalizzato.',
    'Modifica validita: Pulsante per aggiornare le date dal/al del listino.',
]
for item in list_features:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. MODULO AMMINISTRAZIONE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('11. Modulo Amministrazione', level=1)

doc.add_heading('11.1 Diagramma di Flusso', level=2)
add_image(doc, f'{DIAGRAMS}/06_flusso_admin.png', Inches(6.2))

doc.add_heading('11.2 Gestione Utenti', level=2)
doc.add_paragraph(
    'La pagina di gestione utenti permette all\'amministratore di gestire tutti gli account '
    'del sistema.'
)

doc.add_heading('Funzionalita', level=3)
user_features = [
    'Creazione utente: Form modale con campi UserID, Login, Password (obbligatoria), Nome, Tipo utente (dropdown), Email, Codice cliente (ItemID), Codice sede (ItemIDSede).',
    'Modifica utente: Stesso form con campi pre-popolati. La password diventa opzionale (lasciare vuoto per non modificare).',
    'Eliminazione utente: Con conferma preventiva.',
    'Ricerca: Campo di ricerca per login, nome o email.',
    'Export Excel: Scarica l\'elenco completo utenti in formato Excel.',
    'Import Excel: Upload di file Excel/CSV per inserimento massivo utenti. Due modalita: Merge (aggiorna esistenti, aggiunge nuovi) o Replace (sostituisce tutti).',
]
for item in user_features:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('11.3 Import/Export Dati Massivo', level=2)
doc.add_paragraph(
    'La pagina Import permette la gestione massiva dei dati master del sistema '
    'attraverso 4 sezioni (tab):'
)

doc.add_heading('11.3.1 Sezioni Disponibili', level=3)
add_styled_table(doc,
    ['Tab', 'Entita', 'Colonne Chiave', 'Note'],
    [
        ['Prodotti', 'Prodotti + Traduzioni', 'PrdCod, PrdDes, CatCod, FamCod, UM + traduzioni IT/EN/FR/DE', 'Include traduzioni multilingua'],
        ['Prezzi', 'Listini Prezzi', 'PrdCod, LstCod, Prezzo, UM', 'Prezzi per listino'],
        ['Clienti', 'Clienti', 'ItemID, Nome, P.IVA, Indirizzo', 'Dati anagrafici clienti'],
        ['Punti Vendita', 'Punti di Vendita', 'ItemID, ItemIDSede, Nome, Indirizzo, LstCod', 'Sedi per cliente'],
    ],
    [2, 3, 6, 4]
)

doc.add_heading('11.3.2 Processo di Import', level=3)
doc.add_paragraph(
    '1. L\'utente seleziona il tab desiderato.\n'
    '2. Clicca "Import" o trascina un file nell\'area di drop.\n'
    '3. Seleziona la modalita: Merge (aggiorna/inserisci) o Replace (elimina tutti e re-inserisci).\n'
    '4. Clicca "Importa" per avviare il processo.\n'
    '5. Il sistema mostra i risultati: numero di record importati, totale, eventuali errori.\n'
    '6. In caso di errori, viene mostrata una tabella dettagliata con: riga, campo, messaggio di errore.'
)

doc.add_heading('11.3.3 Visualizzazione Dati', level=3)
doc.add_paragraph(
    'Ogni tab include una tabella paginata (10 righe per pagina) per visualizzare i dati '
    'correnti con ricerca testuale. Le card statistiche in alto mostrano i conteggi per '
    'ogni entita e sono cliccabili per navigare al tab corrispondente.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 12. FUNZIONALITA TRASVERSALI
# ═══════════════════════════════════════════════════════════════
doc.add_heading('12. Funzionalita Trasversali', level=1)

doc.add_heading('12.1 Internazionalizzazione (i18n)', level=2)
doc.add_paragraph(
    'L\'applicazione supporta 4 lingue con cambio in tempo reale senza refresh della pagina:'
)
add_styled_table(doc,
    ['Lingua', 'Codice', 'Flag', 'Copertura'],
    [
        ['Italiano', 'it', 'IT', 'Lingua di riferimento (100%)'],
        ['English', 'en', 'EN', 'Traduzione completa'],
        ['Francais', 'fr', 'FR', 'Traduzione completa'],
        ['Deutsch', 'de', 'DE', 'Traduzione completa'],
    ],
    [3, 2, 2, 8]
)
doc.add_paragraph(
    '\nTutte le stringhe dell\'interfaccia passano attraverso il sistema i18next. '
    'La formattazione di numeri, valute e date usa l\'API Intl del browser per adattarsi '
    'automaticamente al locale selezionato. I prodotti hanno traduzioni dedicate (nome, '
    'descrizione) memorizzate nel database.'
)

doc.add_heading('12.2 Accessibilita (WCAG 2.1 AA)', level=2)
doc.add_paragraph('L\'applicazione implementa le seguenti caratteristiche di accessibilita:')
a11y = [
    'Skip Links: Link di salto al contenuto principale, visibili al focus da tastiera.',
    'Focus Trap: Nelle modali, il focus resta confinato all\'interno della modale.',
    'Attributi ARIA: aria-label, aria-hidden, aria-expanded, aria-invalid su tutti gli elementi interattivi.',
    'Icone decorative: Tutte le icone decorative hanno aria-hidden="true".',
    'Tabelle: Uso di scope="col" sugli header delle tabelle.',
    'Contrasto colori: Minimo rapporto 4.5:1 per testo su sfondo (primary-700 o superiore).',
    'Configuratore accessibile: Componente testuale alternativo per la visualizzazione 3D.',
    'Lingua documento: Attributo lang su <html> sincronizzato con la lingua corrente.',
]
for item in a11y:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('12.3 Progressive Web App (PWA)', level=2)
doc.add_paragraph(
    'L\'applicazione e configurata come Progressive Web App completa tramite vite-plugin-pwa 0.21.0 '
    'con Workbox per il caching e un Service Worker registrato con autoUpdate.'
)

doc.add_heading('12.3.1 Web App Manifest', level=3)
doc.add_paragraph(
    'Il manifest viene generato automaticamente da vite-plugin-pwa con i seguenti parametri:'
)
add_styled_table(doc,
    ['Parametro', 'Valore'],
    [
        ['name', 'Portale Recinzioni'],
        ['short_name', 'Recinzioni'],
        ['description', 'Configuratore recinzioni con preventivi e ordini'],
        ['display', 'standalone'],
        ['orientation', 'any'],
        ['theme_color', '#2563eb (blu primary)'],
        ['background_color', '#ffffff'],
        ['start_url / scope', '/'],
        ['lang', 'it'],
        ['categories', 'business, productivity'],
        ['Icona 192x192', 'logo-192.png (PNG)'],
        ['Icona 512x512', 'logo-512.png (PNG + maskable)'],
    ],
    [5, 11]
)

doc.add_heading('12.3.2 Service Worker e Registrazione', level=3)
doc.add_paragraph(
    'Il Service Worker viene registrato in main.jsx tramite registerSW() da virtual:pwa-register. '
    'La configurazione prevede:'
)
sw_features = [
    'registerType: autoUpdate - Il SW si aggiorna automaticamente quando rileva una nuova versione.',
    'skipWaiting: true - Il nuovo SW si attiva immediatamente senza attendere la chiusura delle tab.',
    'clientsClaim: true - Il nuovo SW prende il controllo di tutti i client immediatamente.',
    'cleanupOutdatedCaches: true - Le cache vecchie vengono rimosse automaticamente.',
    'navigateFallback: /index.html - Tutte le navigation request non-API servono index.html (SPA).',
    'navigateFallbackDenylist: /api/, /swagger - Le richieste API e Swagger non vengono intercettate.',
    'Controllo aggiornamenti: Ogni 5 minuti (setInterval) + ad ogni cambio pagina (Layout.jsx).',
]
for item in sw_features:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('12.3.3 Strategie di Caching Runtime', level=3)
doc.add_paragraph(
    'Il Service Worker implementa 6 strategie di caching differenziate per tipo di risorsa:'
)
add_styled_table(doc,
    ['Cache Name', 'Strategia', 'URL Pattern', 'Scadenza', 'Descrizione'],
    [
        ['products-cache', 'StaleWhileRevalidate', '/api/prodotti', '100 entries, 1h', 'Mostra cache, aggiorna in background'],
        ['api-critical-cache', 'NetworkFirst (5s timeout)', '/api/ordini, /api/auth', '50 entries, 5min', 'Rete prima, fallback cache'],
        ['pricelist-cache', 'StaleWhileRevalidate', '/api/listini', '50 entries, 30min', 'Cache con aggiornamento background'],
        ['(health)', 'NetworkOnly', '/api/health', '-', 'Sempre dalla rete'],
        ['images-cache', 'CacheFirst', '.png/.jpg/.svg/.webp', '50 entries, 30gg', 'Immagini statiche'],
        ['fonts-cache', 'CacheFirst', '.woff/.woff2/.ttf', '10 entries, 1 anno', 'Web fonts'],
    ],
    [2.5, 2.5, 3, 2.5, 4]
)

doc.add_heading('12.3.4 Componenti PWA Frontend', level=3)
pwa_components = [
    'InstallPrompt.jsx: Banner di installazione (blu, in basso) che intercetta l\'evento beforeinstallprompt del browser. Presenta pulsanti "Installa" e "Non ora". Se dismissato, non riappare per 7 giorni (localStorage). Si nasconde automaticamente dopo l\'evento appinstalled.',
    'UpdatePrompt.jsx: Banner di aggiornamento (ambra, in alto) che ascolta l\'evento custom pwa-update-available emesso da main.jsx. Pulsanti "Aggiorna ora" (chiama window.__PWA_UPDATE_SW(true)) e "Piu tardi". Fallback con window.location.reload() dopo 500ms.',
    'OfflineBanner.jsx: Banner rosso sottile mostrato automaticamente quando la connessione cade. Usa l\'hook useOnlineStatus che monitora gli eventi online/offline del browser. Si nasconde quando la connessione ritorna.',
    'OfflinePage.jsx: Pagina di fallback (route /offline) con icona WiFi spento, messaggio localizzato e pulsante "Riprova" che ricarica la pagina. Raggiungibile quando l\'utente accede a risorse non in cache senza connessione.',
    'ChunkErrorBoundary (App.jsx): Error Boundary React che intercetta errori di caricamento chunk JavaScript obsoleti (ChunkLoadError, Failed to fetch dynamically imported module). Forza un reload dalla rete con protezione anti-loop (max 1 reload ogni 10 secondi tramite sessionStorage).',
    'useOnlineStatus hook: Hook custom che monitora navigator.onLine e gli eventi online/offline del browser. Restituisce un booleano reattivo usato da OfflineBanner.',
]
for item in pwa_components:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('12.3.5 Meta Tags HTML per PWA', level=3)
doc.add_paragraph(
    'Il file index.html include i meta tag necessari per il supporto PWA su tutti i browser:'
)
meta_items = [
    'theme-color: #2563eb (colore barra del browser)',
    'apple-mobile-web-app-capable: yes (supporto iOS standalone)',
    'apple-mobile-web-app-status-bar-style: black-translucent (stile barra di stato iOS)',
    'apple-mobile-web-app-title: Recinzioni (titolo su iOS)',
    'apple-touch-icon: /apple-touch-icon.png (icona 180x180 per iOS)',
    'Asset in public/: favicon.ico, favicon.svg, apple-touch-icon.png, logo-192.png, logo-512.png',
]
for item in meta_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('12.3.6 Traduzioni PWA', level=3)
doc.add_paragraph(
    'Tutte le 12 stringhe PWA sono localizzate in 4 lingue (IT, EN, FR, DE) nei file i18n sotto '
    'la chiave "pwa". Le chiavi coprono: installTitle, installMessage, installButton, dismissButton, '
    'offlineTitle, offlineMessage, offlineRetry, offlineLimited, updateAvailable, updateMessage, '
    'updateButton, updateDismiss.'
)

doc.add_heading('12.4 Notifiche', level=2)
doc.add_paragraph(
    'Il sistema di notifiche utilizza react-hot-toast per feedback visuale posizionato in alto '
    'a destra con durata di 4 secondi. I toast hanno 4 varianti: successo (verde), errore (rosso), '
    'informazione (blu), caricamento (con spinner).'
)

doc.add_heading('12.5 Generazione PDF', level=2)
doc.add_paragraph(
    'I documenti PDF vengono generati server-side tramite QuestPDF e includono:'
)
pdf_features = [
    'Header: Logo aziendale, dati punto vendita, numero e data ordine.',
    'Sezione indirizzi: Dati fatturazione e consegna.',
    'Tabella prodotti: Con righe padre (articoli) e righe figlio (componenti configurazione).',
    'Totali: Subtotale, IVA 22%, totale complessivo.',
    'Note e condizioni di pagamento (se presenti).',
]
for item in pdf_features:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 13. SPECIFICHE API
# ═══════════════════════════════════════════════════════════════
doc.add_heading('13. Specifiche API REST', level=1)

doc.add_paragraph(
    'Tutti gli endpoint sono esposti sotto il prefisso /api/ e documentati tramite Swagger/OpenAPI '
    '(accessibile in ambiente Development a /swagger).'
)

api_sections = [
    ('Health', [
        ['GET', '/api/health', 'Anonimo', 'Verifica stato sistema e connettivita DB'],
    ]),
    ('Autenticazione', [
        ['POST', '/api/auth/login', 'Anonimo', 'Login con username/password, ritorna JWT + refresh'],
        ['POST', '/api/auth/refresh-token', 'Anonimo', 'Rinnovo JWT tramite refresh token'],
        ['POST', '/api/auth/forgot-password', 'Anonimo', 'Richiesta reset password via email'],
        ['GET', '/api/auth/validate-reset-token', 'Anonimo', 'Validazione token di reset'],
        ['POST', '/api/auth/reset-password', 'Anonimo', 'Impostazione nuova password con token'],
        ['POST', '/api/auth/change-password', 'Autenticato', 'Cambio password (richiede password attuale)'],
    ]),
    ('Prodotti', [
        ['GET', '/api/prodotti', 'Autenticato', 'Lista prodotti paginata con filtri e prezzi'],
        ['GET', '/api/prodotti/{prdCod}', 'Autenticato', 'Dettaglio prodotto con traduzioni'],
        ['GET', '/api/prodotti/filtri', 'Autenticato', 'Categorie, famiglie, gruppi disponibili'],
        ['GET', '/api/prodotti/filtri/famiglie', 'Autenticato', 'Famiglie per categoria'],
    ]),
    ('Configuratore', [
        ['POST', '/api/configuratore/distinta-base', 'Autenticato', 'Calcolo BOM con componenti e prezzi'],
    ]),
    ('Ordini', [
        ['GET', '/api/ordini', 'Autenticato', 'Lista ordini/preventivi con filtri'],
        ['GET', '/api/ordini/{ordNum}', 'Autenticato', 'Dettaglio ordine completo'],
        ['GET', '/api/ordini/dati-anagrafica', 'Autenticato', 'Dati pre-compilazione form preventivo'],
        ['POST', '/api/ordini', 'Autenticato', 'Creazione nuovo preventivo/ordine'],
        ['PUT', '/api/ordini/{ordNum}/conferma', 'Autenticato', 'Conferma preventivo a ordine'],
        ['PUT', '/api/ordini/{ordNum}/invia', 'Autenticato', 'Invio ordine a fornitore via EDI + email'],
        ['DELETE', '/api/ordini/{ordNum}', 'Autenticato', 'Eliminazione ordine'],
        ['DELETE', '/api/ordini/{ordNum}/righe/{rigaNum}', 'Autenticato', 'Eliminazione riga ordine'],
        ['GET', '/api/ordini/{ordNum}/pdf', 'Autenticato', 'Download PDF ordine'],
        ['GET', '/api/ordini/{ordNum}/euritmo/preview', 'Autenticato', 'Preview file EURITMO'],
        ['GET', '/api/ordini/{ordNum}/euritmo', 'Autenticato', 'Download file EURITMO .edi'],
    ]),
    ('Listini', [
        ['GET', '/api/listini/puntivendita', 'Super User', 'Lista PdV con stato listino'],
        ['GET', '/api/listini/{lstCod}', 'Super User', 'Lista prezzi di un listino'],
        ['POST', '/api/listini/crea/{itemIDSede}', 'Super User', 'Crea listino personalizzato'],
        ['GET', '/api/listini/confronto/{itemIDSede}', 'Super User', 'Confronto pubblico vs personalizzato'],
        ['PUT', '/api/listini/{lstCod}/prezzo', 'Super User', 'Aggiorna singolo prezzo'],
        ['PUT', '/api/listini/{lstCod}/validita', 'Super User', 'Aggiorna date validita'],
        ['POST', '/api/listini/{lstCod}/reset', 'Super User', 'Reset a prezzi pubblici'],
        ['DELETE', '/api/listini/{lstCod}', 'Super User', 'Elimina listino personalizzato'],
    ]),
    ('Utenti', [
        ['GET', '/api/utenti', 'Admin', 'Lista utenti con ricerca'],
        ['POST', '/api/utenti', 'Admin', 'Creazione utente'],
        ['PUT', '/api/utenti/{userID}', 'Admin', 'Modifica utente'],
        ['DELETE', '/api/utenti/{userID}', 'Admin', 'Eliminazione utente'],
        ['GET', '/api/utenti/export', 'Admin', 'Export utenti in Excel'],
        ['POST', '/api/utenti/import', 'Admin', 'Import utenti da Excel/CSV'],
    ]),
    ('Import/Export', [
        ['GET', '/api/import/counts', 'Admin', 'Conteggi per tab (prodotti, prezzi, clienti, PdV)'],
        ['GET', '/api/import/{tab}', 'Admin', 'Dati paginati per tab'],
        ['POST', '/api/import/{tab}', 'Admin', 'Import dati da Excel/CSV (merge o replace)'],
        ['GET', '/api/import/{tab}/export', 'Admin', 'Export dati in Excel'],
    ]),
]

for section_name, endpoints in api_sections:
    doc.add_heading(f'13.{api_sections.index((section_name, endpoints))+1} {section_name}', level=2)
    add_styled_table(doc,
        ['Metodo', 'Endpoint', 'Auth', 'Descrizione'],
        endpoints,
        [2, 5, 2.5, 6]
    )
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 14. REQUISITI NON FUNZIONALI
# ═══════════════════════════════════════════════════════════════
doc.add_heading('14. Requisiti Non Funzionali', level=1)

doc.add_heading('14.1 Performance', level=2)
perf = [
    'Lazy loading: Tutte le pagine sono caricate on-demand tramite React.lazy() per ridurre il bundle iniziale.',
    'Code splitting: Vite genera chunk separati per vendor (React, Three.js, i18next) e pagine.',
    'Paginazione server-side: Tutti i dati tabulari sono paginati lato server per evitare download massivi.',
    'Caching PWA: Le risorse statiche sono cachate dal Service Worker per accesso rapido.',
]
for item in perf:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('14.2 Sicurezza', level=2)
sec = [
    'HTTPS obbligatorio in produzione (richiesto per PWA Service Worker).',
    'JWT con scadenza limitata + refresh token per rinnovo automatico.',
    'Protezione brute-force con blocco account dopo 5 tentativi.',
    'Password hashate con BCrypt (cost 12).',
    'CORS configurato con origini specifiche.',
    'Header di sicurezza HTTP (X-Frame-Options, X-Content-Type-Options, CSP) via web.config.',
    'Validazione input server-side su tutti gli endpoint.',
]
for item in sec:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('14.3 Deployment', level=2)
deploy = [
    'Server: Windows Server con IIS.',
    'Build: Script automatizzati (deploy.ps1 per Windows, deploy.sh per Linux).',
    'Process: Build frontend (npm run build) + Publish backend (dotnet publish) + Copia file + Restart IIS App Pool.',
    'Health Check: GET /api/health per verifica post-deploy.',
    'Compressione: Gzip/Brotli abilitati via web.config per risorse statiche.',
    'URL Rewrite: Regole IIS per SPA routing (tutte le route non-API servono index.html).',
]
for item in deploy:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('14.4 Manutenibilita', level=2)
maint = [
    'Codice organizzato per responsabilita (Controller, Services, Models, DTOs).',
    'Frontend modulare con componenti riutilizzabili e hooks custom.',
    'Tool di verifica i18n (check-i18n.js) per garantire completezza traduzioni.',
    'Script di migrazione database idempotenti (IF NOT EXISTS).',
    'Swagger/OpenAPI per documentazione API automatica.',
]
for item in maint:
    doc.add_paragraph(item, style='List Bullet')

# ═══════════════════════════════════════════════════════════════
# Footer su tutte le pagine
# ═══════════════════════════════════════════════════════════════
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Recinzioni Portal - Analisi Funzionale v1.0 - Trafilplast S.r.l. - Marzo 2026')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

# ═══════════════════════════════════════════════════════════════
# SALVA
# ═══════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"Documento salvato: {OUT}")
print(f"Pagine stimate: ~45-50")
